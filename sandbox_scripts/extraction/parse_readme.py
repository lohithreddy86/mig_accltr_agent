#!/usr/bin/env python3
"""
parse_readme.py
===============
Sandbox script: extract structured signals from the user's README.docx.
Runs inside the Podman container.

Called by: Analysis Agent step A0
Input  (--args JSON): {"readme_path": "/source/README.docx"}
Output (stdout JSON): ReadmeSignals dict + quality_score

Mandatory fields (fails with _parse_error if missing):
  - Section 1: Source Database Engine, Dialect Family
  - Section 4: Input Catalog, Input Schema

If the README does not exist or cannot be parsed, returns a minimal
signals dict with quality_score=0.0.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _cell(row, idx: int) -> str:
    try:
        return row.cells[idx].text.strip()
    except Exception:
        return ""


def _table_as_kv(table) -> dict[str, str]:
    """Read a 2-column table as {label: value}. Case-insensitive, strips mandatory markers."""
    result = {}
    for row in table.rows:
        raw_label = _cell(row, 0)
        # Use only first line of multi-line cells (subtitles are for humans, not parser)
        label = raw_label.split("\n")[0] if "\n" in raw_label else raw_label
        # Strip ALL trailing mandatory markers (* and spaces), even repeated ones
        label = label.rstrip(" *").strip()
        # Normalize: collapse whitespace
        label = re.sub(r'\s+', ' ', label)
        value = _cell(row, 1)
        if label:
            result[label] = value
    return result


def _table_as_rows(table) -> list[list[str]]:
    """Read a multi-column table as list of row-lists."""
    rows = []
    for row in table.rows:
        rows.append([c.text.strip() for c in row.cells])
    return rows


# Quality scoring weights — must sum to 1.0
QUALITY_WEIGHTS = {
    "source_engine":    0.30,
    "custom_functions": 0.25,
    "structural_hints": 0.20,
    "problematic_list": 0.10,
    "table_mapping":    0.15,
}
MIN_FUNCTIONS_FOR_SCORE = 3


def compute_quality_score(signals: dict) -> float:
    score = 0.0
    if signals.get("source_engine"):
        score += QUALITY_WEIGHTS["source_engine"]
    if len(signals.get("custom_functions", [])) >= MIN_FUNCTIONS_FOR_SCORE:
        score += QUALITY_WEIGHTS["custom_functions"]
    if signals.get("exception_keyword") and signals.get("transaction_begin"):
        score += QUALITY_WEIGHTS["structural_hints"]
    if len(signals.get("problematic_constructs", [])) > 0:
        score += QUALITY_WEIGHTS["problematic_list"]
    if signals.get("table_mapping"):
        score += QUALITY_WEIGHTS["table_mapping"]
    return round(score, 2)


# ---------------------------------------------------------------------------
# Label matchers (README fields we look for in Section tables)
# ---------------------------------------------------------------------------

SEC1_MAPPINGS = {
    "Source / Application Name":   "app_name",
    "Source Database Engine":      "source_engine",
    "Dialect Family":              "declared_dialect",
    "Procedural Language":         "declared_dialect",  # fallback
    "Business Domain":             "business_domain",
    "Source System Owner / Team":  "team",
    "SME Contact":                 "sme_contact",
    "Desired Output Format":       "output_format",
    "Source Type":                  "source_type",       # stored_procedures | sql_queries
}

SEC2_MAPPINGS = {
    "Proc Naming Convention":      "proc_naming",
    "Are Procs Wrapped in Packages": "has_packages",
    "Script Delimiter / Terminator": "script_delimiter",
    "Exception Syntax Keyword":    "exception_keyword",
    "Transaction Begin Pattern":   "transaction_begin",
    "Commit Pattern":              "transaction_commit",
    "Rollback Pattern":            "transaction_rollback",
    "Dynamic SQL Keyword":         "dynamic_sql_keyword",
    "Cursor Style":                "cursor_style",
    "Bulk Operations Used":        "has_bulk_ops",
}

SEC4_MAPPINGS = {
    "Source Schema(s) Used":       "source_schemas_raw",
    "Temp Table Pattern":          "temp_table_pattern",
    "Are Schema Names Hardcoded":  "schema_hardcoded",
    "Input Catalog":               "input_catalog",
    "Input Schema":                "input_schema",
    "Output Catalog":              "output_catalog",
    "Output Schema":               "output_schema",
}

SEC7_MAPPINGS = {
    "DB Links / Remote DB Calls":  "has_db_links_raw",
    "External System Calls":       "has_external_calls_raw",
    "File I/O Operations":         "has_file_io_raw",
    "Email / Notification Calls":  "has_email_raw",
    "Scheduler / Job Calls":       "has_scheduler_raw",
}


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_readme(readme_path: str) -> dict:
    try:
        import docx
    except ImportError:
        return _empty_signals("python-docx not available")

    path = Path(readme_path)
    if not path.exists():
        return _empty_signals(f"README not found: {readme_path}")

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        return _empty_signals(f"Could not open README: {e}")

    signals: dict = {
        "source_engine": "", "declared_dialect": "", "business_domain": "",
        "sme_contact": "", "proc_start_example": "", "proc_end_example": "",
        "exception_keyword": "", "transaction_begin": "", "transaction_commit": "",
        "transaction_rollback": "", "script_delimiter": "", "has_packages": False,
        "dynamic_sql_keyword": "", "custom_functions": [], "problematic_constructs": [],
        "udf_implementations": [],
        "source_schemas": [], "temp_table_pattern": "", "table_mapping": {},
        "input_catalog": "", "input_schema": [], "output_catalog": "", "output_schema": [],
        "output_format": "AUTO",  # AUTO, TRINO_SQL, or PYSPARK
        "source_type": "stored_procedures",  # stored_procedures | sql_queries
        "proc_io_tables": {},  # {proc_name: {input_tables: [...], output_table: "..."}}
        "procs_to_skip": [],
        "has_db_links": False, "has_external_calls": False,
        "has_file_io": False, "has_scheduler_calls": False,
        "quality_score": 0.0, "_parse_error": None,
    }

    tables = doc.tables
    paragraphs = doc.paragraphs

    # ---------------------------------------------------------------------------
    # Extract UDF code blocks (single-cell tables containing def/function)
    # ---------------------------------------------------------------------------
    _extract_udf_code_blocks(tables, signals)

    # ---------------------------------------------------------------------------
    # Parse all 2-column KV tables
    # ---------------------------------------------------------------------------
    for table in tables:
        if len(table.columns) < 2:
            continue
        kv = _table_as_kv(table)
        # Build case-insensitive lookup: {lowered_label: value}
        kv_lower = {k.lower(): v for k, v in kv.items()}

        for label, field in SEC1_MAPPINGS.items():
            val = kv_lower.get(label.lower())
            if val:
                signals[field] = val

        for label, field in SEC2_MAPPINGS.items():
            val = kv_lower.get(label.lower())
            if val:
                if field in ("has_packages",):
                    signals[field] = "yes" in val.lower()
                else:
                    signals[field] = val

        for label, field in SEC4_MAPPINGS.items():
            val = kv_lower.get(label.lower())
            if val:
                signals[field] = val

        for label, field in SEC7_MAPPINGS.items():
            val = kv_lower.get(label.lower())
            if val:
                v = val.lower()
                key = label.lower()
                is_positive = v.startswith("yes") or (
                    "none" not in v and "no " not in v and len(v) > 3
                )
                if "db links" in key:
                    signals["has_db_links"] = is_positive
                elif "external" in key:
                    signals["has_external_calls"] = is_positive
                elif "file" in key:
                    signals["has_file_io"] = is_positive
                elif "scheduler" in key:
                    signals["has_scheduler_calls"] = is_positive

    # ---------------------------------------------------------------------------
    # Parse multi-column tables (Section 3.1, 3.2, 3.3, 4.2, 6.1, 7.1)
    # ---------------------------------------------------------------------------
    for table in tables:
        rows = _table_as_rows(table)
        if not rows:
            continue

        header = [c.lower().strip() for c in rows[0]]

        # Section 3.1 — Custom Functions
        if _matches_header(header, ["function name", "what it does"]):
            for row in rows[1:]:
                if len(row) >= 2 and row[0] and not _is_example_row(row[0]):
                    signals["custom_functions"].append({
                        "name":               row[0],
                        "what_it_does":       row[1] if len(row) > 1 else "",
                        "example":            row[2] if len(row) > 2 else "",
                        "suggested_equivalent": row[3] if len(row) > 3 else "",
                    })

        # Section 3.3 — Problematic constructs
        elif _matches_header(header, ["construct", "why"]):
            for row in rows[1:]:
                if len(row) >= 2 and row[0] and not _is_example_row(row[0]):
                    signals["problematic_constructs"].append({
                        "construct":          row[0],
                        "why_tricky":         row[1] if len(row) > 1 else "",
                        "frequency":          row[2] if len(row) > 2 else "",
                        "suggested_handling": row[3] if len(row) > 3 else "",
                    })

        # Section 4.2 — Table mapping
        elif _matches_header(header, ["source table", "trino"]):
            for row in rows[1:]:
                if len(row) >= 2 and row[0] and not _is_example_row(row[0]):
                    signals["table_mapping"][row[0]] = row[1]

        # Section 6 — Proc I/O tables (input tables + output table per proc)
        elif _matches_header(header, ["proc name", "input table"]):
            for row in rows[1:]:
                if len(row) >= 2 and row[0] and not _is_example_row(row[0]):
                    signals["proc_io_tables"][row[0]] = {
                        "input_tables": [t.strip() for t in row[1].split(",") if t.strip()],
                        "output_table": row[2].strip() if len(row) > 2 and row[2].strip() else "",
                        "notes":        row[3] if len(row) > 3 else "",
                    }

        # Section 7.1 — Procs to skip
        elif _matches_header(header, ["proc name", "reason"]):
            for row in rows[1:]:
                if len(row) >= 1 and row[0] and not _is_example_row(row[0]):
                    signals["procs_to_skip"].append(row[0])

    # ---------------------------------------------------------------------------
    # Post-process schemas list
    # ---------------------------------------------------------------------------
    raw_schemas = signals.pop("source_schemas_raw", "")
    if raw_schemas:
        signals["source_schemas"] = [s.strip() for s in raw_schemas.split(",") if s.strip()]

    # ---------------------------------------------------------------------------
    # Normalize values for case/space insensitivity downstream
    # ---------------------------------------------------------------------------
    # Table mapping: uppercase keys so source-code names (usually UPPER) match
    if signals["table_mapping"]:
        signals["table_mapping"] = {
            k.strip().upper(): v.strip()
            for k, v in signals["table_mapping"].items()
        }

    # Catalog: strip whitespace, lowercase (single value)
    for cat_field in ("input_catalog", "output_catalog"):
        val = signals.get(cat_field, "")
        if val:
            signals[cat_field] = val.strip().lower()

    # Schema: split comma-separated, strip, lowercase (supports multiple schemas)
    for schema_field in ("input_schema", "output_schema"):
        val = signals.get(schema_field, "")
        if isinstance(val, str) and val:
            signals[schema_field] = [
                s.strip().lower() for s in val.split(",") if s.strip()
            ]

    # Proc I/O tables: strip keys and table names
    if signals["proc_io_tables"]:
        normalized = {}
        for proc, info in signals["proc_io_tables"].items():
            normalized[proc.strip()] = {
                "input_tables": [t.strip() for t in info.get("input_tables", [])],
                "output_table": info.get("output_table", "").strip(),
                "notes": info.get("notes", ""),
            }
        signals["proc_io_tables"] = normalized

    # Output format: normalize to uppercase, validate against known values
    raw_format = signals.get("output_format", "AUTO").strip().upper()
    # Map common aliases
    _FORMAT_ALIASES = {
        "SQL": "TRINO_SQL", "TRINO": "TRINO_SQL", "TRINO SQL": "TRINO_SQL",
        "PYSPARK": "PYSPARK", "SPARK": "PYSPARK", "PYTHON": "PYSPARK",
        "AUTO": "AUTO", "": "AUTO",
    }
    signals["output_format"] = _FORMAT_ALIASES.get(raw_format, raw_format)
    if signals["output_format"] not in ("AUTO", "TRINO_SQL", "PYSPARK"):
        signals["output_format"] = "AUTO"

    # Source type: normalize to stored_procedures | sql_queries
    raw_source_type = signals.get("source_type", "stored_procedures").strip().lower()
    _SOURCE_TYPE_ALIASES = {
        "stored_procedures": "stored_procedures", "stored procedures": "stored_procedures",
        "procedures": "stored_procedures", "procs": "stored_procedures", "sp": "stored_procedures",
        "sql_queries": "sql_queries", "sql queries": "sql_queries",
        "queries": "sql_queries", "sql": "sql_queries", "query": "sql_queries",
        "views": "sql_queries", "reports": "sql_queries", "obiee": "sql_queries",
        "": "stored_procedures",
    }
    signals["source_type"] = _SOURCE_TYPE_ALIASES.get(raw_source_type, "stored_procedures")

    # ---------------------------------------------------------------------------
    # Validate mandatory fields
    # ---------------------------------------------------------------------------
    MANDATORY_FIELDS = {
        "source_engine":    "Section 1 → Source Database Engine (e.g. Oracle 19c, DB2 11.5)",
        "declared_dialect": "Section 1 → Dialect Family (e.g. PL/SQL, T-SQL, SQL PL)",
        "input_catalog":    "Section 4 → Input Catalog (e.g. lakehouse)",
    }

    missing = []
    for field, description in MANDATORY_FIELDS.items():
        val = signals.get(field, "")
        if not val or not val.strip():
            missing.append(f"  - {description}")

    # input_schema is a list — check non-empty
    if not signals.get("input_schema"):
        missing.append("  - Section 4 → Input Schema (e.g. core_banking, risk_dw)")

    if missing:
        error_msg = (
            "README is missing required fields:\n"
            + "\n".join(missing)
            + "\n\nFill in the marked (*) fields in the README template and re-run."
        )
        # Write error to stdout as JSON with _parse_error so the agent gets a clear message
        signals["_parse_error"] = error_msg
        signals["_missing_mandatory"] = missing
        signals["quality_score"] = 0.0
        return signals

    # ---------------------------------------------------------------------------
    # Compute quality score
    # ---------------------------------------------------------------------------
    signals["quality_score"] = compute_quality_score(signals)

    return signals


def _extract_udf_code_blocks(tables, signals: dict) -> None:
    """
    Section 3.2 — Pre-written UDF implementations.
    Each UDF lives in a single-cell table containing Python/PySpark code.
    We identify them by: single column, single row, content contains 'def '.
    """
    for table in tables:
        rows = _table_as_rows(table)
        if len(table.columns) == 1 and rows:
            content = rows[0][0] if rows[0] else ""
            # Heuristic: UDF code block if it contains a function definition
            if re.search(r'\bdef\s+\w+\s*\(', content) and len(content) > 30:
                # Extract function name from first 'def' line
                fn_match = re.search(r'\bdef\s+(\w+)\s*\(', content)
                fn_name = fn_match.group(1) if fn_match else "unknown"
                if not content.startswith("("):  # Skip placeholder rows
                    signals["udf_implementations"].append({
                        "name": fn_name,
                        "code": content,
                    })


def _matches_header(header: list[str], keywords: list[str]) -> bool:
    return all(any(kw in cell for cell in header) for kw in keywords)


def _is_example_row(value: str) -> bool:
    """Skip example/placeholder rows."""
    lower = value.lower()
    return lower.startswith("e.g") or lower.startswith("(example") or lower == "—"


def _empty_signals(reason: str) -> dict:
    return {
        "source_engine": "", "declared_dialect": "", "business_domain": "",
        "sme_contact": "", "proc_start_example": "", "proc_end_example": "",
        "exception_keyword": "", "transaction_begin": "", "transaction_commit": "",
        "transaction_rollback": "", "script_delimiter": "", "has_packages": False,
        "dynamic_sql_keyword": "", "custom_functions": [], "problematic_constructs": [],
        "udf_implementations": [],
        "source_schemas": [], "temp_table_pattern": "", "table_mapping": {},
        "input_catalog": "", "input_schema": [], "output_catalog": "", "output_schema": [],
        "output_format": "AUTO",
        "source_type": "stored_procedures",
        "proc_io_tables": {},
        "procs_to_skip": [],
        "has_db_links": False, "has_external_calls": False,
        "has_file_io": False, "has_scheduler_calls": False,
        "quality_score": 0.0, "_parse_error": reason,
    }


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--args", required=True)
    args = parser.parse_args()
    params = json.loads(args.args)

    signals = parse_readme(params["readme_path"])
    print(json.dumps(signals))


if __name__ == "__main__":
    main()